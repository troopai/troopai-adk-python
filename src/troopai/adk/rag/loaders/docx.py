"""DOCX loader backed by python-docx.

Extracts paragraph and table text from a Word document as a single span (Word
has no intrinsic page model), leaving the chunker to split on paragraph
boundaries. python-docx is an optional dependency; install the ``rag-docx``
extra.
"""

from __future__ import annotations

import asyncio
import logging
from typing import ClassVar, override

from troopai.adk.exceptions.exceptions import DocumentLoadError
from troopai.adk.rag.document import LoadedDocument
from troopai.adk.rag.loaders.base import DocumentLoader

logger = logging.getLogger(__name__)


def _extract_docx(source: str) -> str:
    """Extract paragraph and table text from a .docx (runs in a worker thread)."""
    import docx  # pyright: ignore[reportMissingImports]  # optional dependency (python-docx), guarded at construction

    try:
        document = docx.Document(source)
    except FileNotFoundError as exc:
        raise DocumentLoadError(source, f"DOCX not found: {source}") from exc
    except (ValueError, KeyError, OSError) as exc:
        raise DocumentLoadError(source, f"Could not parse DOCX {source}: {exc}") from exc
    lines = [paragraph.text for paragraph in document.paragraphs if len(paragraph.text.strip()) > 0]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if len(cell.text.strip()) > 0]
            if len(cells) > 0:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


class DOCXLoader(DocumentLoader):
    """Loads a Word (``.docx``) file as a single text document (via python-docx)."""

    requires_packages: ClassVar[tuple[str, ...]] = ("docx",)
    install_extra: ClassVar[str] = "rag-docx"

    def __init__(self) -> None:
        """Verify python-docx is importable, failing fast if the extra is missing.

        Raises:
            ImportError: If python-docx is not installed.
        """
        self.ensure_dependencies()

    @override
    async def load(self, source: str) -> list[LoadedDocument]:
        """Extract text from ``source`` as a single document.

        Args:
            source: Path to a ``.docx`` file.

        Returns:
            A single-element list, or an empty list if the document has no
            extractable text.

        Raises:
            DocumentLoadError: If the document cannot be opened or parsed.
        """
        content = await asyncio.to_thread(_extract_docx, source)
        if len(content.strip()) == 0:
            logger.debug("DOCXLoader: %s has no extractable text", source)
            return []
        return [LoadedDocument(content=content, source=source)]
